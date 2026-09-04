#!/usr/bin/env python3
"""docx2md.py —— 把已排版的笔记 .docx 无损转成本库的 markdown + 内嵌 HTML 表格

保留：粗体 → <strong>，红色(C00000) → <em>，合并单元格 → colspan/rowspan，
      单元格内换行 → <br>，单行单列表 → ### 小节标题。

用法：python3 docx2md.py <in.docx> [--h4 正则] > out.md
      --h4 指定哪些段落作为 #### 小标题（默认 "^\\d+(\\.\\d+)?[.、]?\\s" 或 "^[A-G]-\\d+"）
"""
import sys, re, io
from docx import Document

PIN = []  # 固定 lxml 代理引用，防止 id() 复用导致合并检测误判


def esc(s):
    return s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


def runs_html(p):
    out = []
    for r in p.runs:
        t = r.text
        if not t:
            continue
        col = None
        try:
            if r.font.color is not None and r.font.color.type is not None:
                col = str(r.font.color.rgb)
        except Exception:
            pass
        t = esc(t)
        if col == 'C00000':
            out.append('<em>%s</em>' % t)
        elif r.bold:
            out.append('<strong>%s</strong>' % t)
        else:
            out.append(t)
    return ''.join(out)


def cell_html(c):
    parts = [runs_html(p) for p in c.paragraphs]
    return '<br>'.join(x for x in parts if x.strip())


def table_md(tb):
    rows = tb.rows
    if not rows:
        return ''
    n = len(rows[0].cells)
    grid = []
    for row in rows:
        seen, items = {}, []
        for c in row.cells:
            tc = c._tc
            PIN.append(tc)
            k = id(tc)
            if k in seen:
                items[seen[k]]['colspan'] += 1
            else:
                seen[k] = len(items)
                items.append({'k': k, 'colspan': 1, 'text': cell_html(c)})
        grid.append(items)
    first = {}
    for ri, items in enumerate(grid):
        for it in items:
            first.setdefault(it['k'], ri)
    for ri, items in enumerate(grid):
        for it in items:
            it['skip'] = first[it['k']] != ri
            it['rowspan'] = (sum(1 for rj in range(ri, len(grid))
                                 if any(x['k'] == it['k'] for x in grid[rj]))
                             if not it['skip'] else 1)
    # 首行若是整行合并的引言，则它作 premise，真正的表头顺延到第二行
    lead = (len(grid) > 1 and len(grid[0]) == 1 and grid[0][0]['colspan'] == n)
    hdr_ri = 1 if lead else 0
    out = ['<table class="ftn">']
    for ri, items in enumerate(grid):
        vis = [x for x in items if not x['skip']]
        full = (len(items) == 1 and items[0]['colspan'] == n)
        cls = (' class="premise"' if (lead and ri == 0)
               else ' class="hdr"' if ri == hdr_ri
               else ' class="premise"' if (full and ri == hdr_ri + 1)
               else ' class="note"' if full else '')
        cells = []
        for it in vis:
            tag = 'th' if ri == hdr_ri else 'td'
            a = ''
            if it['colspan'] > 1:
                a += ' colspan="%d"' % it['colspan']
            if it['rowspan'] > 1:
                a += ' rowspan="%d"' % it['rowspan']
            cells.append('<%s%s>%s</%s>' % (tag, a, it['text'], tag))
        out.append('<tr%s>%s</tr>' % (cls, ''.join(cells)))
    out.append('</table>')
    s = '\n'.join(out)
    return re.sub(r'<th([^>]*)><strong>(.*?)</strong></th>', r'<th\1>\2</th>', s)


def convert(path, h4pat=r'^(\d+(\.\d+)?[.、]?\s|[A-G]-\d+)', drop=()):
    d = Document(path)
    body, pi, ti = [], 0, 0
    for el in d.element.body:
        tag = el.tag.split('}')[1]
        if tag == 'p':
            if pi < len(d.paragraphs):
                t = d.paragraphs[pi].text.strip()
                if t and t not in drop:
                    if re.match(r'^[一二三四五六七八九十]+、', t):
                        body.append('### ' + re.sub(r'^([一二三四五六七八九十]+、)\s*(\d[\d.]*\s*)?', r'\1', t))
                    elif re.match(h4pat, t):
                        body.append('#### ' + re.sub(r'\s+', '　', t))
                    else:
                        body.append(t)
            pi += 1
        elif tag == 'tbl':
            tb = d.tables[ti]
            ti += 1
            if len(tb.rows) == 1 and len(tb.rows[0].cells) == 1:
                body.append('### ' + tb.rows[0].cells[0].text.strip())
            else:
                body.append(table_md(tb))
    return '\n\n'.join(body)


if __name__ == '__main__':
    src = sys.argv[1]
    drop = tuple(a for a in sys.argv[2:] if not a.startswith('--'))
    sys.stdout.write(convert(src, drop=drop))
